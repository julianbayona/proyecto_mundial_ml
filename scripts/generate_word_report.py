from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "outputs" / "documents"
DOCX_PATH = OUTPUT_DIR / "informe_final_el_factor_local.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE_GRAY = "E8EEF5"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text_run = paragraph.add_run("Actualizar campo en Word si es necesario.")
    text_run.italic = True

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(fld_end)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "El Factor Local - Proyecto final de Machine Learning"
    header.style = styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("El Factor Local")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(
        "Analisis del impacto de la sede y las condiciones del torneo en los resultados de la Copa Mundial de Futbol"
    )
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Proyecto final de Machine Learning")
    run.bold = True
    run.font.size = Pt(13)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2600, 6760])
    rows = [
        ("Asignatura", "Machine Learning"),
        ("Alcance", "Copa Mundial de la FIFA 1930-2014"),
        ("Enfoque supervisado", "Regresion logistica para predecir avance de fase"),
        ("Enfoque no supervisado", "K-Means y clustering jerarquico por edicion"),
        ("Fecha de elaboracion", "2026-05-22"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE_GRAY)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Documento consolidado a partir del desarrollo por fases, notebooks, metricas, figuras y modelos generados en el proyecto."
    ).italic = True
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            if len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(value)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_figure(doc: Document, image_name: str, caption: str, width: float = 6.2) -> None:
    path = PROJECT_ROOT / "outputs" / "figures" / image_name
    if not path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[Imagen sugerida no encontrada: {image_name}] {caption}").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


def add_toc(doc: Document) -> None:
    add_heading(doc, "Tabla de contenido", 1)
    sections = [
        "1. Introduccion",
        "2. Objetivos",
        "3. Problematica",
        "4. Requerimientos funcionales",
        "5. Documentacion del dataset",
        "6. Preparacion y calidad de los datos",
        "7. Desarrollo de la metodologia",
        "8. Modelo supervisado: Regresion logistica",
        "9. Modelo no supervisado: K-Means y clustering jerarquico",
        "10. Errores y decisiones tecnicas",
        "11. Analisis critico del desempeno",
        "12. Descripcion de la solucion y analisis de datos",
        "13. Reflexion sobre implementacion en entorno real",
        "14. Conclusiones y recomendaciones",
        "15. Referencias",
        "16. Anexos",
    ]
    for item in sections:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run(
        "Nota: si se desea una tabla con numeros de pagina automaticos, se puede insertar o actualizar una tabla de contenido nativa desde Microsoft Word."
    ).italic = True
    doc.add_page_break()


def load_json(relative_path: str) -> dict:
    with (PROJECT_ROOT / relative_path).open("r", encoding="utf-8") as file:
        return json.load(file)


def build_report() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title_page(doc)
    add_toc(doc)

    supervised_metrics = load_json("outputs/metrics/supervised_metrics.json")
    clustering_metrics = load_json("outputs/metrics/clustering_metrics.json")
    final_summary = load_json("outputs/metrics/final_summary.json")
    clean_df = pd.read_csv(PROJECT_ROOT / "data" / "processed" / "world_cup_matches_clean.csv")
    supervised_df = pd.read_csv(PROJECT_ROOT / "data" / "processed" / "supervised_dataset.csv")
    clusters_df = pd.read_csv(PROJECT_ROOT / "data" / "processed" / "world_cup_editions_clusters.csv")

    add_heading(doc, "1. Introduccion", 1)
    add_paragraph(
        doc,
        "La Copa Mundial de Futbol es un evento deportivo con una historia extensa, formatos cambiantes y condiciones contextuales que pueden influir en el rendimiento de las selecciones. Este proyecto analiza si variables como el continente sede, la experiencia historica y la confederacion del equipo aportan informacion para explicar o predecir el avance mas alla de la fase de grupos.",
    )
    add_paragraph(
        doc,
        "El trabajo combina aprendizaje supervisado y no supervisado. En la parte supervisada se entrena una regresion logistica para predecir avance de fase. En la parte no supervisada se agrupan ediciones mundialistas mediante K-Means y clustering jerarquico.",
    )

    add_heading(doc, "2. Objetivos", 1)
    add_heading(doc, "2.1 Objetivo general", 2)
    add_paragraph(
        doc,
        "Analizar el impacto de condiciones contextuales de la Copa Mundial en los resultados de las selecciones participantes, aplicando algoritmos supervisados y no supervisados de Machine Learning sobre datos historicos del torneo.",
    )
    add_heading(doc, "2.2 Objetivos especificos", 2)
    add_numbered(
        doc,
        [
            "Recopilar, limpiar y documentar datos historicos de la Copa Mundial.",
            "Construir un dataset supervisado sin data leakage a nivel seleccion-edicion.",
            "Entrenar un baseline y una regresion logistica para predecir avance de fase.",
            "Agrupar ediciones mundialistas con K-Means y clustering jerarquico.",
            "Evaluar los modelos con metricas estandar e interpretar sus resultados.",
            "Documentar limitaciones, errores, decisiones tecnicas y posibles usos reales.",
        ],
    )

    add_heading(doc, "3. Problematica", 1)
    add_paragraph(
        doc,
        "La pregunta central es si las condiciones de sede y el rendimiento historico disponible antes del torneo ayudan a explicar el avance de una seleccion mas alla de la fase de grupos. El problema es complejo porque el futbol depende de factores tecnicos, tacticos, fisicos, historicos y contextuales que no siempre estan disponibles en los datasets.",
    )
    add_paragraph(
        doc,
        "Ademas, la Copa Mundial cambio de formato varias veces entre 1930 y 2014. Esto significa que la variable de avance no representa exactamente el mismo contexto competitivo en todas las epocas.",
    )

    add_heading(doc, "4. Requerimientos funcionales", 1)
    add_table(
        doc,
        ["Req.", "Descripcion", "Evidencia"],
        [
            ["RF01", "Cargar datasets crudos desde data/raw.", "src/data_loader.py"],
            ["RF02", "Cruzar resultados con etapas del Mundial.", "world_cup_matches_raw.csv"],
            ["RF03", "Limpiar nulos, duplicados y valores atipicos.", "src/preprocessing.py"],
            ["RF04", "Generar visualizaciones de EDA.", "outputs/figures/"],
            ["RF05", "Construir dataset supervisado sin leakage.", "supervised_dataset.csv"],
            ["RF06", "Entrenar baseline y modelo supervisado.", "supervised_metrics.json"],
            ["RF07", "Entrenar modelos de clustering.", "clustering_metrics.json"],
            ["RF08", "Serializar modelos y consolidar metricas.", "outputs/models/"],
            ["RF09", "Ejecutar pipeline completo desde main.py.", "final_summary.json"],
        ],
        [850, 5200, 3310],
    )

    add_heading(doc, "5. Documentacion del dataset", 1)
    add_heading(doc, "5.1 Origen", 2)
    add_bullets(
        doc,
        [
            "International Football Results from 1872 to 2024: fuente principal de resultados de partidos internacionales.",
            "FIFA World Cup Dataset: fuente complementaria con informacion de etapas del torneo mediante WorldCupMatches.csv.",
            "WorldCups.csv: archivo de informacion por edicion, disponible en data/raw aunque no fue el insumo principal de modelado.",
        ],
    )
    add_heading(doc, "5.2 Alcance", 2)
    add_paragraph(
        doc,
        "El alcance final fue 1930-2014. Se excluyeron 2018, 2022 y 2026 porque la fuente complementaria disponible no incluia la columna stage para esas ediciones.",
    )
    add_table(
        doc,
        ["Dataset procesado", "Filas", "Columnas"],
        [
            ["world_cup_matches_raw.csv", "836", "11"],
            ["world_cup_matches_clean.csv", "836", "15"],
            ["supervised_dataset.csv", "424", "9"],
            ["world_cup_editions_clusters.csv", "20", "24"],
        ],
        [4300, 1600, 1600],
    )
    add_heading(doc, "5.3 Variables principales", 2)
    add_table(
        doc,
        ["Variable", "Significado", "Uso"],
        [
            ["year", "Anio de la edicion mundialista.", "Identificador temporal"],
            ["home_team / away_team", "Selecciones que disputan el partido.", "Construccion de partidos y equipos"],
            ["home_score / away_score", "Goles del partido.", "EDA y construccion historica, no feature actual"],
            ["stage", "Fase del torneo.", "Identificar grupos y rondas posteriores"],
            ["is_group_stage", "Indicador de fase de grupos o equivalente.", "Target y agregaciones"],
            ["host_continent", "Continente sede del Mundial.", "Feature contextual"],
            ["advanced_group_stage", "Target binario de avance.", "Modelo supervisado"],
            ["historical_appearances", "Participaciones previas.", "Feature pre-torneo"],
            ["historical_win_rate", "Tasa historica de victorias previas.", "Feature pre-torneo"],
            ["historical_avg_goals_scored", "Promedio historico de goles previos.", "Feature pre-torneo"],
        ],
        [2200, 4700, 2460],
    )

    add_heading(doc, "6. Preparacion y calidad de los datos", 1)
    add_heading(doc, "6.1 Carga y cruce", 2)
    add_paragraph(
        doc,
        "El cruce final entre fuentes logro cobertura completa para el alcance 1930-2014. Se aplico normalizacion de nombres historicos y se permitio cruce con orientacion invertida de local/visitante cuando las fuentes diferian.",
    )
    add_table(
        doc,
        ["Indicador", "Resultado"],
        [
            ["Partidos de Copa Mundial hasta 2014", "836"],
            ["Partidos excluidos por alcance", "200"],
            ["Partidos sin stage asignado", "0"],
            ["Cobertura del join", "100.00%"],
        ],
        [3900, 5460],
    )
    add_heading(doc, "6.2 Limpieza", 2)
    add_bullets(
        doc,
        [
            "No se encontraron nulos en columnas criticas.",
            "No se encontraron duplicados exactos ni duplicados por identidad completa del partido.",
            "Se conservaron rematches historicos que comparten equipos y anio pero tienen fecha o marcador distinto.",
            "No se detectaron marcadores negativos ni valores extremos superiores a 20 goles por equipo.",
            "Se agregaron is_group_stage, host_continent, goal_difference y total_goals.",
        ],
    )
    add_figure(doc, "distribucion_goles_raw.png", "Figura 1. Distribucion de goles locales y visitantes.")
    add_figure(doc, "partidos_por_edicion.png", "Figura 2. Cantidad de partidos por edicion mundialista.")
    add_figure(doc, "goles_promedio_por_edicion.png", "Figura 3. Promedio de goles por partido por edicion.")

    add_heading(doc, "7. Desarrollo de la metodologia", 1)
    add_numbered(
        doc,
        [
            "Fase 0: configuracion del proyecto, carpetas, requirements y modulos base.",
            "Fase 1: carga y cruce de datos para obtener stage.",
            "Fase 2: limpieza de datos y generacion del dataset limpio.",
            "Fase 3: analisis exploratorio y visualizaciones.",
            "Fase 4: construccion de features supervisadas sin data leakage.",
            "Fase 5: entrenamiento y evaluacion del modelo supervisado.",
            "Fase 6: clustering de ediciones mundialistas.",
            "Fase 7: integracion, README, notebook final y verificacion de outputs.",
        ],
    )
    add_figure(doc, "rendimiento_por_sede.png", "Figura 4. Tasa de avance segun relacion con continente sede.")
    add_figure(doc, "balance_clases.png", "Figura 5. Balance de clases del target supervisado.")
    add_figure(doc, "top_selecciones_avance.png", "Figura 6. Selecciones con mayor cantidad de avances.")
    add_figure(doc, "matriz_correlacion.png", "Figura 7. Matriz de correlacion de variables numericas de partidos.")

    add_heading(doc, "8. Modelo supervisado: Regresion logistica", 1)
    add_heading(doc, "8.1 Justificacion del algoritmo", 2)
    add_paragraph(
        doc,
        "La regresion logistica fue seleccionada porque es interpretable, adecuada para clasificacion binaria y permite analizar el peso relativo de variables contextuales e historicas. Tambien es apropiada para un proyecto academico con un dataset pequeno.",
    )
    add_heading(doc, "8.2 Separacion de datos", 2)
    add_paragraph(
        doc,
        "Se uso train_test_split con test_size=0.2, random_state=42 y stratify=y. La division fue de 339 registros para entrenamiento y 85 para prueba. La validacion de hiperparametros se reviso de forma exploratoria con validacion cruzada sobre el entrenamiento, pero el modelo final mantuvo C=1.0 por estabilidad y coherencia con el plan.",
    )
    add_table(
        doc,
        ["Particion", "Registros", "Proporcion clase 0", "Proporcion clase 1"],
        [["Entrenamiento", "339", "0.531", "0.469"], ["Prueba", "85", "0.5294", "0.4706"]],
        [2300, 1700, 2600, 2760],
    )
    add_heading(doc, "8.3 Metricas", 2)
    add_table(
        doc,
        ["Modelo", "Accuracy", "Precision", "Recall", "F1", "AUC-ROC"],
        [
            [
                "Baseline",
                str(supervised_metrics["baseline_accuracy"]),
                str(supervised_metrics["baseline_precision"]),
                str(supervised_metrics["baseline_recall"]),
                str(supervised_metrics["baseline_f1_score"]),
                str(supervised_metrics["baseline_auc_roc"]),
            ],
            [
                "Regresion logistica",
                str(supervised_metrics["model_accuracy"]),
                str(supervised_metrics["precision"]),
                str(supervised_metrics["recall"]),
                str(supervised_metrics["f1_score"]),
                str(supervised_metrics["auc_roc"]),
            ],
        ],
        [2600, 1350, 1350, 1350, 1350, 1360],
    )
    add_paragraph(
        doc,
        "El modelo supera al baseline en accuracy, F1 y AUC. Aun asi, el desempeno es moderado y debe interpretarse con cautela porque la accuracy no alcanzo la referencia orientativa de 0.65.",
    )
    add_figure(doc, "matrices_confusion.png", "Figura 8. Matrices de confusion para baseline y regresion logistica.")
    add_figure(doc, "curva_roc.png", "Figura 9. Curva ROC del modelo supervisado.")
    add_figure(doc, "coeficientes_modelo.png", "Figura 10. Coeficientes de la regresion logistica.")

    add_heading(doc, "9. Modelo no supervisado: K-Means y clustering jerarquico", 1)
    add_heading(doc, "9.1 Justificacion de algoritmos", 2)
    add_paragraph(
        doc,
        "K-Means permite evaluar diferentes valores de k mediante inercia y coeficiente de silueta. El clustering jerarquico es especialmente util en datasets pequenos, como este caso de 20 ediciones, porque permite visualizar relaciones mediante un dendrograma.",
    )
    add_heading(doc, "9.2 Variables de clustering", 2)
    add_paragraph(
        doc,
        "Se construyeron 19 variables agregadas por edicion: goles promedio, diferencia media de goles, numero de partidos, numero de equipos, proporcion de fase de grupos, tasa de avance general, proporcion de equipos de region sede, representacion por confederacion y tasas de avance por confederacion.",
    )
    add_heading(doc, "9.3 Metricas", 2)
    add_table(
        doc,
        ["Indicador", "Resultado"],
        [
            ["Numero de ediciones", str(clustering_metrics["n_editions"])],
            ["Mejor k", str(clustering_metrics["best_k"])],
            ["Mejor silueta", str(clustering_metrics["best_silhouette"])],
            ["Varianza PCA componente 1", str(clustering_metrics["pca_explained_variance_ratio"][0])],
            ["Varianza PCA componente 2", str(clustering_metrics["pca_explained_variance_ratio"][1])],
        ],
        [4200, 5160],
    )
    add_figure(doc, "kmeans_codo.png", "Figura 11. Metodo del codo para K-Means.")
    add_figure(doc, "kmeans_silueta.png", "Figura 12. Coeficiente de silueta por k.")
    add_figure(doc, "dendrograma.png", "Figura 13. Dendrograma del clustering jerarquico.")
    add_figure(doc, "clusters_pca.png", "Figura 14. Visualizacion PCA de clusters.")
    add_heading(doc, "9.4 Interpretacion de clusters", 2)
    add_table(
        doc,
        ["Cluster", "Ediciones", "Interpretacion"],
        [
            ["0", "1934-1982 excepto 1950", "Ediciones intermedias o de formato clasico/transicional."],
            ["1", "1986, 1990, 1994, 1998, 2002, 2006, 2010, 2014", "Ediciones modernas con mas equipos, mas partidos y menor promedio de goles."],
            ["2", "1930, 1950", "Ediciones tempranas y atipicas con pocos equipos y alta diferencia media de goles."],
        ],
        [1200, 3900, 4260],
    )

    add_heading(doc, "10. Errores y decisiones tecnicas", 1)
    add_bullets(
        doc,
        [
            "El primer join tuvo baja cobertura porque WorldCupMatches.csv solo llegaba hasta 2014.",
            "Se decidio limitar el alcance a 1930-2014 para no inventar etapas.",
            "Se encontraron diferencias de nombres como USA y United States, Germany FR y Germany, Korea Republic y South Korea.",
            "Se preservaron rematches historicos en lugar de eliminarlos como duplicados.",
            "La instalacion completa de requirements tuvo problemas con Jupyter y rutas largas en Windows; se instalaron dependencias minimas para modelado.",
            "Durante el clustering aparecio una advertencia no bloqueante de joblib sobre deteccion de nucleos fisicos.",
        ],
    )

    add_heading(doc, "11. Analisis critico del desempeno", 1)
    add_paragraph(
        doc,
        "El desempeno del modelo supervisado es mejor que el baseline, pero moderado. Esto sugiere que las variables historicas y contextuales contienen senal predictiva, aunque no explican completamente el avance deportivo.",
    )
    add_bullets(
        doc,
        [
            "Posibles errores: falsos positivos y falsos negativos derivados de selecciones con poca historia previa o cambios de formato.",
            "Posibles sesgos: sobrerrepresentacion de selecciones historicamente fuertes y confederaciones con mas cupos.",
            "Sobreajuste: se redujo mediante un modelo simple, train/test estratificado y uso de pocas features pre-torneo.",
            "Limitacion metodologica: no se uso un conjunto de validacion independiente por el tamano pequeno del dataset; se hizo revision exploratoria con validacion cruzada.",
        ],
    )

    add_heading(doc, "12. Descripcion de la solucion y analisis de datos", 1)
    add_paragraph(
        doc,
        "La solucion final consiste en un pipeline reproducible ejecutable desde main.py. El pipeline carga datos, limpia, genera features, entrena modelos, produce figuras, guarda metricas y consolida un resumen final. La estructura modular permite reutilizar cada fase por separado y revisar resultados desde notebooks.",
    )
    add_paragraph(
        doc,
        f"El archivo final_summary.json confirma que todos los outputs esperados estan presentes: {final_summary['all_expected_outputs_present']}.",
    )

    add_heading(doc, "13. Reflexion sobre implementacion en entorno real", 1)
    add_paragraph(
        doc,
        "En un entorno real, el modelo podria servir como modulo exploratorio para estimar probabilidades pre-torneo de avance de fase. Sin embargo, no deberia usarse como predictor definitivo sin ampliar las fuentes de datos.",
    )
    add_bullets(
        doc,
        [
            "Agregar rankings FIFA o Elo antes de cada torneo.",
            "Incluir datos de plantillas, edad, experiencia internacional y lesiones.",
            "Actualizar el dataset con 2018 y 2022 usando una fuente complementaria con stage.",
            "Usar validacion temporal, entrenando con torneos pasados y probando en torneos posteriores.",
            "Exponer el modelo como API solo despues de monitorear calibracion, sesgos y drift de datos.",
        ],
    )

    add_heading(doc, "14. Conclusiones y recomendaciones", 1)
    add_paragraph(
        doc,
        "El proyecto cumplio su objetivo academico y tecnico. Se construyo un pipeline completo desde datos crudos hasta modelos supervisados y no supervisados, con documentacion y outputs reproducibles.",
    )
    add_bullets(
        doc,
        [
            "La ventaja regional aparece como una senal descriptiva, pero no permite afirmar causalidad.",
            "La regresion logistica supera al baseline, aunque con desempeno moderado.",
            "La experiencia historica y algunas confederaciones tienen asociaciones positivas con el avance.",
            "El clustering separa ediciones tempranas, intermedias y modernas de forma interpretable.",
            "Los resultados deben presentarse como evidencia exploratoria y no como prediccion definitiva.",
        ],
    )

    add_heading(doc, "15. Referencias", 1)
    add_bullets(
        doc,
        [
            "Kaggle. International Football Results from 1872 to 2024.",
            "Kaggle. FIFA World Cup Dataset: WorldCupMatches.csv y WorldCups.csv.",
            "scikit-learn developers. Documentacion de LogisticRegression, DummyClassifier, Pipeline, KMeans, AgglomerativeClustering y metricas de evaluacion.",
            "SciPy documentation. scipy.cluster.hierarchy para dendrogramas y linkage jerarquico.",
            "pandas documentation. Herramientas de carga, limpieza, merge y agregacion de datos.",
            "Matplotlib y Seaborn documentation. Visualizacion estadistica en Python.",
            "Documentacion interna del proyecto: archivos Markdown en la carpeta docs/ y notebooks por fase.",
        ],
    )

    add_heading(doc, "16. Anexos", 1)
    add_heading(doc, "16.1 Archivos principales generados", 2)
    add_bullets(
        doc,
        [
            "data/processed/world_cup_matches_clean.csv",
            "data/processed/supervised_dataset.csv",
            "data/processed/world_cup_editions_clusters.csv",
            "outputs/metrics/supervised_metrics.json",
            "outputs/metrics/clustering_metrics.json",
            "outputs/metrics/final_summary.json",
            "outputs/models/logistic_regression_model.joblib",
            "outputs/models/kmeans_model.joblib",
            "outputs/models/hierarchical_clustering_model.joblib",
        ],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
